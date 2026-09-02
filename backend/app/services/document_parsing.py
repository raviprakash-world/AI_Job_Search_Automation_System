from pathlib import Path

import pdfplumber
from docx import Document

from app.core.errors import UnsupportedFileTypeError

SUPPORTED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}


def parse_docx(path: Path) -> str:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def parse_pdf(path: Path) -> str:
    text_chunks: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def parse_document(path: Path, mime_type: str) -> str:
    kind = SUPPORTED_MIME_TYPES.get(mime_type)
    if kind == "docx":
        return parse_docx(path)
    if kind == "pdf":
        return parse_pdf(path)
    raise UnsupportedFileTypeError(f"Unsupported file type: {mime_type}")
