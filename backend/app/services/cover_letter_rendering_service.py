from pathlib import Path

from docx import Document


def render_cover_letter_docx(
    *, full_name: str | None, email: str | None, phone: str | None, company_name: str, body_text: str, destination: Path
) -> None:
    """Renders a plain single-column cover letter — no tables/text boxes/images,
    same ATS-safety rationale as resume_rendering_service.py."""
    doc = Document()

    doc.add_heading(full_name or "Cover Letter", level=1)
    contact_parts = [p for p in [email, phone] if p]
    if contact_parts:
        doc.add_paragraph(" | ".join(contact_parts))

    doc.add_paragraph(f"Dear {company_name} Hiring Team,")

    for paragraph in body_text.split("\n\n"):
        stripped = paragraph.strip()
        if stripped:
            doc.add_paragraph(stripped)

    doc.add_paragraph(f"Sincerely,\n{full_name or ''}")

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destination))
