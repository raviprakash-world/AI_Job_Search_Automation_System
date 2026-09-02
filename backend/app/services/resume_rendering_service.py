from pathlib import Path

from docx import Document

from app.schemas.resumes import StructuredResumeContent


def _format_date_range(start, end, is_current: bool) -> str:
    start_str = start.strftime("%b %Y") if start else "?"
    if is_current:
        end_str = "Present"
    else:
        end_str = end.strftime("%b %Y") if end else "?"
    return f"{start_str} – {end_str}"


def render_resume_docx(content: StructuredResumeContent, destination: Path) -> None:
    """Renders structured resume content to an ATS-safe DOCX.

    Deliberately avoids tables, text boxes, columns, and images — a plain
    single-column flow of headings/paragraphs/bullets is the most reliably
    parseable format for ATS systems, and is what the spec asks for.
    """
    doc = Document()

    doc.add_heading(content.full_name or "Resume", level=1)

    contact_parts = [p for p in [content.email, content.phone, content.location] if p]
    contact_parts.extend(f"{label}: {url}" for label, url in content.links.items())
    if contact_parts:
        doc.add_paragraph(" | ".join(contact_parts))

    if content.professional_summary:
        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(content.professional_summary)

    if content.skills:
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(content.skills))

    if content.experiences:
        doc.add_heading("Experience", level=2)
        for exp in content.experiences:
            header = doc.add_paragraph()
            header.add_run(f"{exp.title} — {exp.company}").bold = True
            meta_line = _format_date_range(exp.start_date, exp.end_date, exp.is_current)
            if exp.location:
                meta_line = f"{exp.location} | {meta_line}"
            meta = doc.add_paragraph()
            meta.add_run(meta_line).italic = True
            for bullet in exp.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    if content.education:
        doc.add_heading("Education", level=2)
        for edu in content.education:
            line = edu.institution
            if edu.degree:
                line = f"{edu.degree}" + (f", {edu.field_of_study}" if edu.field_of_study else "") + f" — {edu.institution}"
            doc.add_paragraph(line)
            if edu.start_date or edu.end_date:
                meta = doc.add_paragraph()
                meta.add_run(_format_date_range(edu.start_date, edu.end_date, False)).italic = True

    if content.certifications:
        doc.add_heading("Certifications", level=2)
        for cert in content.certifications:
            label = f"{cert.name} ({cert.issuer})" if cert.issuer else cert.name
            doc.add_paragraph(label, style="List Bullet")

    if content.projects:
        doc.add_heading("Projects", level=2)
        for project in content.projects:
            header = doc.add_paragraph()
            header.add_run(project.name).bold = True
            if project.description:
                doc.add_paragraph(project.description)
            if project.technologies:
                doc.add_paragraph(f"Technologies: {', '.join(project.technologies)}")

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destination))
