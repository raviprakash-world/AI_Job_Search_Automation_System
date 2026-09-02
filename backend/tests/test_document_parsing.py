from pathlib import Path

import pytest
from docx import Document

from app.core.errors import UnsupportedFileTypeError
from app.services.document_parsing import parse_document, parse_docx


def test_parse_docx_extracts_paragraph_text(tmp_path: Path):
    doc = Document()
    doc.add_heading("Jane Doe", level=1)
    doc.add_paragraph("Senior Engineer with 8 years of experience.")
    path = tmp_path / "resume.docx"
    doc.save(str(path))

    text = parse_docx(path)
    assert "Jane Doe" in text
    assert "Senior Engineer with 8 years of experience." in text


def test_parse_docx_extracts_table_content(tmp_path: Path):
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Python"
    table.rows[0].cells[1].text = "Expert"
    path = tmp_path / "resume_table.docx"
    doc.save(str(path))

    text = parse_docx(path)
    assert "Python" in text
    assert "Expert" in text


def test_parse_document_rejects_unsupported_mime_type(tmp_path: Path):
    path = tmp_path / "resume.txt"
    path.write_text("plain text resume")

    with pytest.raises(UnsupportedFileTypeError):
        parse_document(path, "text/plain")


def test_parse_document_dispatches_docx(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("Dispatch check")
    path = tmp_path / "resume.docx"
    doc.save(str(path))

    text = parse_document(
        path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "Dispatch check" in text
